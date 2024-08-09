import os
from fpdf import FPDF
from docx import Document
from PIL import Image, ImageDraw, ImageFont
import textwrap
import streamlit as st
import unicodedata
import base64
import io
import json
import pyperclip
import markdown 
from docx.shared import Pt, RGBColor


def escape_special_chars(text):
    replacements = {
        '<': '&lt;',
        '>': '&gt;',
        '&': '&amp;',
        '"': '&quot;',
        "'": '&#39;',
        '`': '\`'  # Escape backticks
    }
    for search, replace in replacements.items():
        text = text.replace(search, replace)
    return text



class UTF8FPDF(FPDF):
    def __init__(self):
        super().__init__()
        self.add_font('DejaVu', '', '/usr/share/fonts/truetype/dejavu/DejaVuSansCondensed.ttf', uni=True)
        self.add_font('Courier', '', '/usr/share/fonts/truetype/dejavu/DejaVuSansMono.ttf', uni=True)  # Add a monospaced font
    
    def header(self):
        pass
    
    def footer(self):
        pass

def remove_emojis(text):
    return ''.join(c for c in text if not unicodedata.category(c).startswith('So'))


def export_to_pdf(messages):
    pdf = UTF8FPDF()
    pdf.add_page()
    
    pdf.set_font("DejaVu", size=12)
    
    for message in messages:
        role = message['role'].capitalize()
        content = remove_emojis(message['content'])
        
        if '```' in content:
            # Split the content based on code blocks
            parts = content.split('```')
            for i, part in enumerate(parts):
                if i % 2 == 0:  # Normal text
                    pdf.set_font("DejaVu", size=12)
                    pdf.multi_cell(0, 10, f"{role}: {part}")
                else:  # Code block
                    pdf.set_font("Courier", size=12)
                    pdf.multi_cell(0, 10, part)
                pdf.ln(5)
        else:
            pdf.set_font("DejaVu", size=12)
            pdf.multi_cell(0, 10, f"{role}: {content}")
            pdf.ln(5)
    
    pdf_output = io.BytesIO()
    pdf_output_bytes = pdf.output(dest='S').encode('latin1')  # Get the PDF as bytes
    pdf_output.write(pdf_output_bytes)
    pdf_output.seek(0)
    return pdf_output.getvalue()



def export_to_md(messages):
    md_content = "# Chat Export\n\n"
    md_content += "| Role | Content |\n|------|--------|\n"
    for message in messages:
        content = remove_emojis(message['content']).replace("\n", "<br>")
        md_content += f"| {message['role'].capitalize()} | {content} |\n"
    
    return md_content.encode()

def markdown_to_html(markdown_text):
    # Convert Markdown to HTML using the markdown library
    return markdown.markdown(markdown_text)

def export_to_html(messages):
    html_content = """
    <html>
    <head>
        <style>
            table { border-collapse: collapse; width: 100%; }
            th, td { border: 1px solid black; padding: 8px; text-align: left; }
            th { background-color: #f2f2f2; }
            .markdown-table { border-collapse: collapse; width: 100%; }
            .markdown-table th, .markdown-table td { border: 1px solid black; padding: 8px; text-align: left; }
            .markdown-table th { background-color: #f2f2f2; }
            pre { background-color: #f8f8f8; border: 1px solid #ccc; padding: 10px; overflow: auto; }
        </style>
    </head>
    <body>
        <h1>Chat Export</h1>
        <table class="markdown-table">
            <tr><th>Role</th><th>Content</th></tr>
    """
    for message in messages:
        content = remove_emojis(message['content']).replace("\n", "<br>")
        if message['role'] in ['user', 'assistant']:
            # Split the content to handle code blocks separately
            parts = content.split('```')
            content_html = ""
            for i, part in enumerate(parts):
                if i % 2 == 0:
                    # Regular text
                    content_html += markdown_to_html(part)
                else:
                    # Code block
                    content_html += f"<pre>{part}</pre>"
            html_content += f"<tr><td>{message['role'].capitalize()}</td><td>{content_html}</td></tr>"
        else:
            html_content += f"<tr><td>{message['role'].capitalize()}</td><td>{content}</td></tr>"
    html_content += "</table></body></html>"
    return html_content.encode()

def export_to_docx(messages):
    doc = Document()
    doc.add_heading('Chat Export', 0)
    
    # Define 'Code' style
    styles = doc.styles
    code_style = styles.add_style('Code', 1)  # 1 is for character style
    code_style.font.name = 'Courier New'
    code_style.font.size = Pt(10)
    code_style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    code_style.font.bold = False

    for message in messages:
        content = remove_emojis(message['content'])
        parts = content.split('```')
        for i, part in enumerate(parts):
            if i % 2 == 0:
                paragraph = doc.add_paragraph(f"{message['role'].capitalize()}: {part}")
                paragraph.style = 'Normal'
            else:
                paragraph = doc.add_paragraph(part)
                paragraph.style = 'Code'
    
    docx_output = io.BytesIO()
    doc.save(docx_output)
    docx_output.seek(0)
    return docx_output

def export_to_image(messages, format):
    width, height = 800, len(messages) * 100 + 100
    image = Image.new('RGB', (width, height), color='white')
    draw = ImageDraw.Draw(image)
    font_regular = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 14)
    font_code = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSansMono.ttf", 12)
    
    y_text = 10
    for message in messages:
        content = remove_emojis(message['content'])
        parts = content.split('```')
        for i, part in enumerate(parts):
            if i % 2 == 0:
                lines = textwrap.wrap(f"{message['role'].capitalize()}: {part}", width=70)
                for line in lines:
                    draw.text((10, y_text), line, font=font_regular, fill='black')
                    y_text += 20
            else:
                lines = textwrap.wrap(part, width=70)
                for line in lines:
                    draw.text((10, y_text), line, font=font_code, fill='black')
                    y_text += 20
            y_text += 10
    
    img_output = io.BytesIO()
    image.save(img_output, format=format)
    img_output.seek(0)
    return img_output

def export_chat(messages, format='html'):
    if format == 'pdf':
        return export_to_pdf(messages)
    elif format == 'md':
        return export_to_md(messages)
    elif format == 'html':
        return export_to_html(messages)
    elif format == 'docx':
        return export_to_docx(messages)
    elif format in ['png', 'jpeg']:
        return export_to_image(messages, format)
    else:
        raise ValueError(f"Unsupported format: {format}")

def add_export_button(key_suffix=""):
    col1, col2, col3 = st.columns([1, 1, 1])

    with col1:
        export_format = st.selectbox("Format", ['pdf', 'md', 'html', 'docx', 'png', 'jpeg'], key=f"export_format_{key_suffix}")

    with col2:
        if st.button("Export Chat"):
            try:
                exported_content = export_chat(st.session_state.messages, format=export_format)
                st.success(f"Chat exported as {export_format.upper()}!")
                
                if export_format in ['pdf', 'docx', 'png', 'jpeg']:
                    
                    st.download_button(
                        label=f"Download {export_format.upper()}",
                        data=exported_content,
                        file_name=f"chat_export.{export_format}",
                        mime=f"application/{export_format}"
                    )
                elif export_format in ['md', 'html']:
                    b64 = base64.b64encode(exported_content).decode()
                    href = f'<a href="data:text/{export_format};base64,{b64}" download="chat_export.{export_format}">Download {export_format.upper()}</a>'
                    st.markdown(href, unsafe_allow_html=True)
            except Exception as e:
                st.error(f"Export error: {str(e)}")

    with col3:
        if st.button("Copy to Clipboard", key=f"copy_button_{key_suffix}"):
            # Create formatted chat text
            formatted_chat = "\n".join(
                f"{'user' if m['role'] == 'user' else 'assistant'}: {m['content']}"
                for m in st.session_state.messages
            )
            print(formatted_chat)
            pyperclip.copy(formatted_chat)
            st.success("Chat copied to clipboard!")
